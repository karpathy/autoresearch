#!/usr/bin/env python3
"""Build Round 14 Full Research Proposal Word document (source → regenerates .docx)."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT = "/Users/mehmetaltay/Documents/GitHub/autoresearch/Cambridge project/FRP_R14_MAltay.docx"

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "Adaptive Computerized Dynamic Assessment: Synthesizing Eye-Tracking and Reader Profiles "
    "for Inclusive English Reading Comprehension"
)
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run(
    "Full Research Proposal — Cambridge English Funded Research Programme (Round 14)"
).italic = True

sections = [
    (
        "Alignment with the Round 14 call",
        (
            "This proposal addresses Round 14 priorities around emerging constructs and digitally mediated "
            "assessment by modelling reading comprehension as jointly evidenced by (a) product scores, "
            "(b) traceable processing behaviours, and (c) responsiveness to calibrated mediation. Rather than "
            "treating comprehension as a single latent trait inferred only from item responses, the study "
            "articulates how moment-to-moment allocation of attention—indexed through eye-movement parameters—"
            "can inform which scaffolding is triggered for which learner on which item type. The design therefore "
            "targets explanatory fairness: it links observable processing signatures to adaptive support aligned "
            "with the Zone of Proximal Development (ZPD), supporting more equitable inferences for diverse EFL "
            "readers when digital tests risk under-representing emerging abilities (McNamara, 2014; Poehner & Wang, 2021)."
        ),
    ),
    (
        "Background and context of the study",
        (
            "McNamara (2014) argued that language assessment must confront technological change and shifting "
            "communicative realities; digital delivery now makes process capture feasible at scale, reopening "
            "questions about what should count as evidence in reading assessment. Two research programmes are "
            "especially relevant.\n\n"
            "Eye-tracking and process-oriented validity in language testing. Eye-movement recording has become a "
            "standard complement to outcome scores when researchers ask whether tests elicit the cognitive processes "
            "their constructs imply (Bax, 2013). In high-stakes reading assessment, process evidence can support "
            "validity arguments by describing how candidates allocate attention across text and items (Lim, 2020). "
            "Related work shows how individual differences interact with gaze during L2 reading-for-testing tasks "
            "(Tywoniw, 2023). Eye-tracking has also been applied beyond reading-for-comprehension screens—for "
            "example to analyse attention to visual cues in L2 listening tests (Batty, 2020)—illustrating how "
            "process data can be embedded in computer-delivered tasks more broadly. Syntheses and introductions in "
            "second language research position eye-tracking as a method that increasingly pairs with other evidence "
            "sources and with interdisciplinary tooling (Godfroid, Winke, & Conklin, 2020). A recent review maps the "
            "growth of eye-tracking research in L2 assessment and foregrounds cognitive validity and processing "
            "patterns among recurring themes (Cao & Ma, 2025). Across these lines of "
            "work, gaze indices are typically treated as auxiliary evidence that constrains interpretation of "
            "performance rather than as standalone proficiency scores.\n\n"
            "Dynamic assessment and technology-mediated mediation. Sociocultural Dynamic Assessment (DA) integrates "
            "assessment and instruction so that the aim includes estimating emerging abilities through cooperative "
            "mediation (Lantolf & Poehner, 2010; Poehner & Lantolf, 2005). Computerised Dynamic Assessment (C-DA) "
            "scales interventionist mediation through graduated prompts and yields multi-index profiles of "
            "responsiveness (Poehner, Zhang, & Lu, 2015; Zhang & Lu, 2019). Contemporary accounts emphasise DA’s "
            "relevance to digital ecologies and to diagnosing development during mediation (Poehner & Wang, 2021). "
            "Recent classroom-oriented implementations further demonstrate how hybrid computerised DA can combine "
            "interactionist insights with scalable delivery in L2 reading contexts (Jin & Liu, 2024). However, many "
            "C-DA systems still deliver largely fixed mediation ladders and therefore under-use continuous evidence "
            "about how a learner is processing the text at hand.\n\n"
            "Heterogeneity in L2 reading. Large-scale eye-movement corpora underscore that predictors of fluency "
            "and comprehension need not align uniformly across multilingual readers (Kuperman et al., 2023), "
            "motivating profile-sensitive assessment designs. Reader profiles tied to competence and motivational–"
            "cognitive resources provide stable priors for interpreting gaze (Alexander, 2005; Dinsmore et al., 2018).\n\n"
            "Synthesis. The project integrates process-oriented validity logic from eye-tracking research in "
            "language testing with technology-mediated DA: real-time indices of processing difficulty become "
            "inputs to mediation selection, while outcome and mediation traces remain primary evidentiary objects.\n\n"
            "Literature sources informing this synthesis were identified through Web of Science Core Collection "
            "searches (e.g., topic queries pairing computerised dynamic assessment with reading/L2; eye tracking with "
            "language testing) and supplemented with publisher metadata (Crossref) for DOI-verified citations. "
            "Complementary discovery in Google Scholar and ResearchGate was used to confirm bibliographic details for "
            "handbooks and recent CALL articles where needed."
        ),
    ),
    (
        "Research gap",
        (
            "Three gaps motivate Adaptive Interventionist Dynamic Assessment (A-IDA). First, scalable Interventionist "
            "C-DA trades away moment-to-moment responsiveness: mediation is typically pre-sequenced rather than "
            "conditional on ongoing comprehension signals. Second, although eye-tracking has been used to evaluate "
            "tests and tasks, fewer implementations close the loop by using gaze features to adapt scaffolding inside "
            "the assessment session itself. Third, product-only scoring leaves explanatory gaps when learners with "
            "different profiles produce similar scores or when scores understate emerging ability.\n\n"
            "A-IDA targets a synthesis: eye-tracking metrics and reader-profile classifications jointly trigger "
            "individualised scaffolding during computerised reading, enabling empirical comparison against traditional "
            "C-DA and a non-mediated control."
        ),
    ),
    (
        "Research questions",
        (
            "1. Are there significant differences in reading comprehension outcomes and response times among "
            "participants assigned to A-IDA, a traditional computerised Interventionist DA sequence, and a "
            "non-mediated control condition?\n\n"
            "2. How do comprehension outcomes and response times vary across reader profiles (e.g., highly competent, "
            "effortful, challenged) under A-IDA compared with traditional C-DA?\n\n"
            "3. What relationships obtain between eye-movement parameters (e.g., fixation-based difficulty proxies, "
            "regressive saccades) and the type and frequency of scaffolding cues triggered during testing?"
        ),
    ),
    (
        "Research design",
        (
            "The study follows design-based research (DBR) using the ADDIE cycle (Analysis, Design, Development, "
            "Implementation, Evaluation) to build, pilot, and refine A-IDA across iterative micro-cycles. DBR fits "
            "because the artefact (an adaptive testing interface) and its instrumentation co-evolve with evidence "
            "from learner use (McNamara, 2014; Poehner & Wang, 2021).\n\n"
            "Figure 1 summarises the processing architecture. In live operation, gaze streams are segmented into "
            "fixations and saccades; derived features feed a profile-aware decision layer that selects mediation "
            "from a validated bank mapped to item-level skills (e.g., inferencing, cohesion). Insert Figure 1 in Word "
            "as a flowchart (Insert → SmartArt → Process, or an exported diagram).\n\n"
            "Table 1 outlines a 12-month timeline aligned to the programme window (December 2026–December 2027). "
            "Exact months will be adjusted for ethics approval and term dates."
        ),
    ),
]

for heading, body in sections:
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)

# Table 1
doc.add_paragraph()
t = doc.add_table(rows=1, cols=4)
hdr = t.rows[0].cells
hdr[0].text = "Phase (ADDIE)"
hdr[1].text = "Period"
hdr[2].text = "Primary outputs"
hdr[3].text = "Evaluation focus"

rows = [
    ("Analysis", "Dec 2026 – Jan 2027", "Needs analysis; risk register; mediation taxonomy v0", "Construct map ↔ Cambridge item types"),
    ("Design", "Feb – Mar 2027", "Wireframes; cue database v1; pilot battery; analysis preregistration", "Expert review (n = 5)"),
    ("Development", "Apr – Jul 2027", "EyeLink–PsychoPy/PyQt integration; online feature pipeline; rules v1", "Latency; reliability of gaze features"),
    ("Implementation", "Aug – Sep 2027", "Pilot (n ≈ 20–30); tune thresholds", "Usability; cue–item alignment"),
    ("Evaluation", "Oct – Dec 2027", "Main study; analysis; reporting; handover pack", "Effects; process–outcome models"),
]
for phase, period, outputs, ev in rows:
    r = t.add_row().cells
    r[0].text = phase
    r[1].text = period
    r[2].text = outputs
    r[3].text = ev

doc.add_paragraph()
doc.add_paragraph(
    "Figure 1. A-IDA cycle: gaze acquisition → event detection (fixations, regressions) → "
    "profile-aware mediation selection → delivery via the testing UI → logged outcomes for analysis."
)

doc.add_heading("Methodology", level=1)

method_parts = [
    (
        "Rationale for proposed methods",
        (
            "Eye-tracking is included because it captures how comprehension unfolds during test completion, aligning "
            "with process-oriented validity arguments in language testing research (Bax, 2013; Godfroid, Winke, & "
            "Conklin, 2020; Lim, 2020). Fixation- and regression-based features provide proxies for difficulty and "
            "strategic reanalysis within trial windows compatible with mediated items. Profile classification "
            "(Alexander, 2005; Dinsmore et al., 2018) supplies priors that reduce overfitting idiosyncratic gaze noise. "
            "The learning potential score (LPS) family remains appropriate because it indexes responsiveness to "
            "mediation relative to unassisted performance (Poehner et al., 2015; Zhang & Lu, 2019). Scikit-learn "
            "decision rules offer an interpretable first-generation mapping from features to mediation levels—"
            "auditable for stakeholders—before more complex models are considered once trace datasets exist."
        ),
    ),
    (
        "Participants and research team",
        (
            "Participants will be EFL undergraduates in English-medium programmes at a public university in Türkiye. "
            "Vocabulary screening will target intermediate–upper intermediate readers (approximately 5,000–6,000 word "
            "families) to align with Cambridge-level tasks. Reader profiles will follow Alexander (2005) as informed "
            "by interest and prior knowledge (Dinsmore et al., 2018). Target recruitment is approximately 60–90 "
            "participants (20–30 per arm), preceded by a pilot to stabilise mediation thresholds. The team combines "
            "applied linguists/assessment researchers with a software engineer for the adaptive interface—consistent "
            "with principled mediation design in computerised DA (Jin & Liu, 2024)."
        ),
    ),
    (
        "Instruments and data collection",
        (
            "Reading materials and items will be drawn from validated Cambridge English examinations (e.g., B2 First, "
            "C1 Advanced, or IELTS Academic reading), supporting alignment with Cambridge constructs. An expert panel "
            "(five applied linguists) will review text–mediation mapping before pilots.\n\n"
            "Eye-tracking will use the SR Research EyeLink Portable Duo available in the lab. The stack (pylink, "
            "PsychoPy, PyQt5; NumPy/Pandas for online features) prioritises timing fidelity. Mediation banks follow "
            "graduated prompting while allowing profile-conditioned branches (Jin & Liu, 2024; Poehner et al., 2015)."
        ),
    ),
    (
        "Procedure and analysis",
        (
            "Participants will be randomly assigned to A-IDA, standard C-DA, or control. Dependent variables include "
            "accuracy, response time, mediation level, and LPS contrasts. Omnibus comparisons (e.g., one-way ANOVA or "
            "robust alternatives) will contrast arms; moderation by reader profile will use factorial or regression "
            "models; gaze–mediation relations will use correlational/regression frameworks with multiple-comparison "
            "control as appropriate. Effect sizes and confidence intervals will be reported alongside p-values.\n\n"
            "Ethics: institutional approval, informed consent, withdrawal rights, secure storage of logs, protocols "
            "for incidental findings. Data access will follow least privilege and pseudonymised identifiers.\n\n"
            "Open-science practices: the team will preregister the main confirmatory analysis plan (e.g., OSF) before "
            "the main data collection window, specifying primary contrasts, planned contrasts or simple effects for "
            "profiles, and rules for handling outliers and missing gaze segments. Materials that can be shared "
            "without violating Cambridge item copyright will be documented in a replication packet.\n\n"
            "Limitations: laboratory calibration requirements; generalisation initially to similar proficiency bands "
            "and task types; mediation rules versioned as A-IDA v1 for replication. Transparency: mediation codebook, "
            "analysis scripts, and (where copyright permits) aggregated process summaries for Cambridge stakeholders."
        ),
    ),
]

for sub, body in method_parts:
    doc.add_heading(sub, level=2)
    doc.add_paragraph(body)

doc.add_heading("Potential implications for Cambridge University Press & Assessment and the wider field", level=1)
doc.add_paragraph(
    "For Cambridge English, the project offers a research pathway toward digital assessments and learning products "
    "that pair high-quality items with interpretable, profile-sensitive feedback. Because stimuli come from Cambridge "
    "examinations, results can inform how diverse profiles engage with existing tasks—supporting item development, "
    "automated tutoring overlays, and transparency in score meaning when mediation is present.\n\n"
    "For the wider field, the work demonstrates how emerging-construct claims can be supported by linked process–"
    "product evidence (cf. Bax, 2013; Kuperman et al., 2023; Poehner & Wang, 2021). Documenting mediation rules and "
    "gaze-derived features advances auditable inclusive design beyond fixed adaptive algorithms; later work can "
    "explore proxies where eye-tracking hardware is unavailable."
)

doc.add_heading("References", level=1)

refs = [
    "Alexander, P. A. (2005). The path to competence: A lifespan developmental perspective on reading. "
    "Journal of Literacy Research, 37(4), 413–436.",
    "Bax, S. (2013). The cognitive processing of candidates during reading tests: Evidence from eye-tracking. "
    "Language Testing, 30(4), 441–465. https://doi.org/10.1177/0265532212473244",
    "Batty, A. (2020). An eye-tracking study of attention to visual cues in L2 listening tests. "
    "Language Testing, 38(4), 511–535. https://doi.org/10.1177/0265532220951504",
    "Cao, X., & Ma, Z. (2025). The review on eye-tracking studies in L2 assessment. "
    "Colombian Applied Linguistics Journal, 27(2), 51–63. https://doi.org/10.14483/22487085.22043",
    "Dinsmore, D. L., Fox, E., Parkinson, M. M., & Bilgili, F. (2018). The interplay of prior knowledge, interest, "
    "and reading comprehension. Learning and Individual Differences, 65, 477–490.",
    "Godfroid, A., Winke, P., & Conklin, K. (2020). Exploring the depths of second language processing with eye "
    "tracking: An introduction. Second Language Research, 36(3), 243–256. https://doi.org/10.1177/0267658320922578",
    "Jin, C., & Liu, Y. (2024). Diagnosing and promoting learners’ L2 inferential reading development through hybrid "
    "computerised dynamic assessment in the Chinese EFL classroom. Computer Assisted Language Learning. "
    "Advance online publication. https://doi.org/10.1080/09588221.2024.2421521",
    "Kuperman, V., Siegelman, N., Schroeder, S., Acartürk, C., Alexeeva, S., Amenta, S., Bertram, R., Bonandrini, R., "
    "Brysbaert, M., Chernova, D., Da Fonseca, S. M., Dirix, N., Duyck, W., Fella, A., Frost, R., Gattei, C. A., "
    "Kalaitzi, A., Lõo, K., Marelli, M., … Usal, K. A. (2023). Text reading in English as a second language: "
    "Evidence from the Multilingual Eye-Movements Corpus. Studies in Second Language Acquisition, 45(1), 3–37. "
    "https://doi.org/10.1017/S0272263121000954",
    "Lantolf, J. P., & Poehner, M. E. (2004). Dynamic assessment of L2 development: Bringing the past into the future. "
    "Journal of Applied Linguistics, 1(1), 49–72.",
    "Lantolf, J. P., & Poehner, M. E. (2010). Dynamic assessment in the classroom: Vygotskian praxis for second language "
    "development. Language Teaching Research, 15(1), 11–33.",
    "Lim, H. (2020). Exploring the validity evidence of a high-stake, second language reading test: An eye-tracking study. "
    "Language Testing in Asia, 10, Article 14. https://doi.org/10.1186/s40468-020-00107-0",
    "McNamara, T. (2014). 30 years on—Evolution or revolution? Language Assessment Quarterly, 11(2), 226–232.",
    "Poehner, M. E., & Lantolf, J. P. (2005). Dynamic assessment in the language classroom. Language Teaching Research, "
    "9(3), 233–265.",
    "Poehner, M. E., & Wang, Z. (2021). Dynamic Assessment and second language development. Language Teaching, 54(4), "
    "472–490. https://doi.org/10.1017/S0261444820000555",
    "Poehner, M. E., Zhang, J., & Lu, X. (2015). Computerized dynamic assessment (C-DA): Diagnosing L2 development "
    "according to learner responsiveness to mediation. Language Testing, 32(3), 337–357. "
    "https://doi.org/10.1177/0265532214560390",
    "Tywoniw, R. (2023). Compensatory effects of individual differences, language proficiency, and reading behavior: "
    "An eye-tracking study of second language reading assessment. Frontiers in Communication, 8, Article 1176986. "
    "https://doi.org/10.3389/fcomm.2023.1176986",
    "Vygotsky, L. S. (1978). Mind in society: The development of higher psychological processes. Harvard University Press.",
    "Zhang, J., & Lu, X. (2019). Measuring and supporting second language development using computerized dynamic "
    "assessment. Language and Sociocultural Theory, 6(1), 92–115. https://doi.org/10.1558/lst.31710",
]

for r in refs:
    p = doc.add_paragraph(r)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)

doc.save(OUT)
print("Wrote", OUT)
